from airtable import Airtable
from docx import Document
from docx2pdf import convert
import os

TABLE_NAME = 'Table 1'
FOLDER_NAME = TABLE_NAME + ' Submissions'
BASE_KEY = 'appsExXNgItazXQbG'
MY_KEY = ''
fault_count = 0

airtable = Airtable(
    BASE_KEY, TABLE_NAME, MY_KEY
)
submissions = airtable.get_all(sort = 'Division', maxRecords = 10)

try:
    if FOLDER_NAME in os.listdir():
        os.chdir(FOLDER_NAME)
    else:
        os.mkdir(FOLDER_NAME)
        print('Created base folder')
        os.chdir(FOLDER_NAME)

    delete_count = 0
    for division in ['A', 'B','C','D','E','F','G']:
        if division not in os.listdir():
            os.mkdir(division)
        else:
            for file in os.listdir(division):
                os.remove(division + '/' + file)
                delete_count += 1
    print('Deleted previous files:', delete_count)

        
    print('Number of files:' , len(submissions))
    print('Converting data...')
    for test in submissions:  
        try:
            document = Document() #'''This was the issue, create new document each time'''
            file_name = test['fields']['Division'] + '/' + test['fields']['Name'] + '.docx'
            for i in test['fields']:
                p = document.add_heading(str(i)).bold = True
                q = document.add_paragraph(test['fields'][i])
            r = document.add_paragraph(test['createdTime'])

            if test['fields']['Division'] in os.listdir():
                document.save(file_name)
                convert(file_name)
                os.remove(file_name)

        except Exception as e:
            fault_count += 1
            print(e)

except Exception as e:
    print(e)



