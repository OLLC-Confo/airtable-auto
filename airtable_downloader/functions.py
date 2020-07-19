from airtable import Airtable
from docx2pdf import convert
from docx import Document
import os


division_map = {
    'X': ['A','B','C','D','E','F','G'],
    'XI': ['A','B','C','D','E','F']
}

class LectureSubmission:
    
    def __init__(self, base_key, api_key, table_name, path, CLASS):
        self.table_name = table_name
        self.airtable = Airtable(
            base_key, table_name, api_key
        ).get_all(sort = 'Division', maxRecords = 10) #maxRecords = 10
        self.count = len(self.airtable)
        self.path = path
        self.CLASS = CLASS
        self.divisions = division_map[CLASS]

    def create_empty_folders(self):
        os.chdir(self.path)
        print('PATH changed')
        
        FOLDER_NAME = self.table_name
        if not FOLDER_NAME in os.listdir():
            os.mkdir(FOLDER_NAME)
            print('Created base folder.')
        else:
            print('Base folder exists.')
        os.chdir(FOLDER_NAME)
        delete_count = 0
        for division in self.divisions:
            if division not in os.listdir():
                os.mkdir(division)
            else:
                for file in os.listdir(division):
                    os.remove(division + '/' + file)
                    delete_count += 1
        print('Deleted previous files:', delete_count, '\n')

    def create_pdf(self):
        fault_count = 0
        for test in self.airtable:
            try:
                document = Document()
                document.sections[0].footer.paragraphs[0].text = f"\t\t{self.CLASS} {self.table_name} {test['fields']['Name']} {test['createdTime']}"
                file_name = test['fields']['Division'] + '/' + test['fields']['Name'] + '.docx'
                for i in test['fields']:
                    document.add_heading(str(i)).bold = True
                    document.add_paragraph(test['fields'][i])
                document.save(file_name)
                convert(file_name)
                os.remove(file_name)
                yield True
            except Exception as e:
                print(e)
                fault_count += 1
        print("Fault count:", fault_count)