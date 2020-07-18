# from airtable import Airtable

from docx2pdf import convert
from docx import Document
import os

class LectureSubmission:
    
    def __init__(self, base_key, api_key, table_name, path, divisions):
        self.table_name = table_name
        self.airtable = Airtable(
            base_key, table_name, api_key
        ).get_all(sort = 'Division', maxRecords=10)
        self.count = len(self.airtable)
        self.path = path
        self.divisions = divisions


    def create_empty_folders(self):
        os.chdir(self.path)
        print('PATH changed')
        
        FOLDER_NAME = 'Submissions'
        if not FOLDER_NAME in os.listdir():
            os.mkdir(FOLDER_NAME)
            print('Created base folder.')
        else:
            print('Base folder exists.')
        os.chdir(FOLDER_NAME)
        delete_count = 0
        for division in self.divisions:
            if division not in os.listdir():
                os.mkdir(division)
            else:
                for file in os.listdir(division):
                    os.remove(division + '/' + file)
                    delete_count += 1
        print('Deleted previous files:', delete_count, '\n')

    def create_pdf(self):
        fault_count = 0
        for test in self.airtable:
            try:
                document = Document()
                file_name = test['fields']['Division'] + '/' + test['fields']['Name'] + '.docx'
                for i in test['fields']:
                    document.add_heading(str(i)).bold = True
                    document.add_paragraph(test['fields'][i])
                document.add_paragraph(test['createdTime'])
                document.save(file_name)
                convert(file_name)
                os.remove(file_name)
                yield True
            except Exception as e:
                print(e)
                fault_count += 1
        print("Fault count:", fault_count)