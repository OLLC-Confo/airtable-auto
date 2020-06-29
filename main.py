from airtable import Airtable
from docx import Document
from docx2pdf import convert
import os

TABLE_NAME = 'Table 1'
FOLDER_NAME = TABLE_NAME + ' Submissions'
BASE_KEY = 'appsExXNgItazXQbG'
MY_KEY = 'enter your key here'

airtable = Airtable(
    BASE_KEY, TABLE_NAME, MY_KEY
)

try:
    if FOLDER_NAME in os.listdir():
        os.chdir(FOLDER_NAME)
    else:
        os.mkdir(FOLDER_NAME)
        os.chdir(FOLDER_NAME)
    for i in os.listdir():
        os.remove(i)
except Exception as e:
    print(e)

submission = airtable.get_all(maxRecords = 1)

document = Document()

for test in submission:
    file_name = test['fields']['Name'] + '.docx'
    for i in test['fields']:
        p = document.add_heading(str(i)).bold = True
        p = document.add_paragraph(test['fields'][i])
    p = document.add_paragraph(test['createdTime'])

    document.save(file_name)
    convert(file_name)
    os.remove(file_name)